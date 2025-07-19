import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from flask import Flask, request, render_template, send_from_directory, jsonify
from werkzeug.utils import secure_filename
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
REPORT_FOLDER = "reports"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)

@app.route('/about')
def about():
    return render_template('about.html')



def extract_co_mapping(xls, sheet_name, prefix):
    """Extract only the CO-PO or CO-PSO mapping table (stop at PO Attainment)."""
    df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
    df.dropna(axis=1, how='all', inplace=True)

    header_idx = None
    raw_header = None

    # Locate CO-PO header
    for i, row in df.iterrows():
        vals = [str(x).strip() for x in row if pd.notna(x)]
        if len(vals) >= 2 and ((vals[0].upper() == "CO" or vals[0].upper() == f"{prefix}1")
                               and vals[1].upper().startswith(prefix)):
            header_idx = i
            raw_header = list(df.iloc[i])
            break
    if header_idx is None:
        raise ValueError(f"Could not locate CO‑{prefix} header in '{sheet_name}' sheet.")

    # Normalize header
    header = [str(x).strip() for x in raw_header]
    if header[0].upper() == f"{prefix}1":
        header[0] = "CO"

    # Stop when "PO Attainment" row appears
    stop_idx = None
    for i in range(header_idx + 1, len(df)):
        row = df.iloc[i]
        if any(str(x).strip().lower() == "po attainment" for x in row if pd.notna(x)):
            stop_idx = i
            break

    # Extract only valid mapping rows
    if stop_idx:
        data = df.iloc[header_idx + 1 : stop_idx].copy()
    else:
        data = df.iloc[header_idx + 1 :].copy()

    # Filter out empty rows
    mask = ~data.iloc[:, 0].astype(str).str.lower().str.startswith("overall co attainment")
    data = data[mask].dropna(how="all").reset_index(drop=True)
    data = data.iloc[:, : len(header)]
    data.columns = header
    data.fillna("-", inplace=True)

    # Round numeric columns
    for col in data.columns[1:]:
        data[col] = pd.to_numeric(data[col], errors='coerce').round(2).fillna("-")

    return {"heading": f"Mapping of CO with {prefix}s:", "data": data}


def extract_co_pso_mapping(xls, sheet_name="PSO Attainment"):
    """Extracts Mapping of COs with PSOs from 'PSO Attainment' sheet."""
    df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
    df.dropna(axis=1, how='all', inplace=True)

    header_idx = None
    raw_header = None

    for i, row in df.iterrows():
        vals = [str(x).strip() for x in row if pd.notna(x)]
        if len(vals) >= 2 and vals[0].lower().startswith("course outcomes"):
            if any("ps" in x.lower() for x in vals[1:]):
                header_idx = i
                raw_header = list(df.iloc[i])
                break

    if header_idx is None:
        raise ValueError("Could not locate CO‑PSO header in 'PSO Attainment' sheet.")

    header = [str(x).strip() for x in raw_header]
    header[0] = "Course Outcomes"

    # Extract data
    data = df.iloc[header_idx + 1:].copy()

    #  Remove 'Course' rows
    data = data[~data.iloc[:, 0].astype(str).str.strip().str.lower().eq("course")]
    data = data.dropna(how="all").reset_index(drop=True)

    data = data.iloc[:, :len(header)]
    data.columns = header
    data.fillna("-", inplace=True)

    # Round numeric values
    for col in data.columns[1:]:
        data[col] = pd.to_numeric(data[col], errors='coerce').round(2).fillna("-")

    return {"heading": "Mapping of COs with PSOs:", "data": data}

def extract_po_attainment(xls, sheet_name):
    """Extracts the single-row PO Attainment summary table."""
    df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
    df.dropna(axis=1, how='all', inplace=True)

    pa_idx = None
    for i, row in df.iterrows():
        if any(str(x).strip().lower() == "po attainment" for x in row if pd.notna(x)):
            pa_idx = i
            break
    if pa_idx is None:
        return None

    header_idx = None
    for i in range(pa_idx + 1, len(df)):
        row = df.iloc[i].astype(str).str.strip().tolist()
        if row and row[0].lower().startswith("co program"):
            continue
        if row and row[0].upper() == "CO" and any(cell.upper().startswith("PO") for cell in row[1:]):
            header_idx = i
            break
    if header_idx is None:
        return None

    header = df.iloc[header_idx].dropna().tolist()
    data_row = df.iloc[header_idx + 1].fillna("-").tolist()

    # round numeric values
    rounded_row = []
    for val in data_row[: len(header)]:
        try:
            num = round(float(val), 2)
            rounded_row.append(num)
        except:
            rounded_row.append(val)

    po_df = pd.DataFrame([rounded_row], columns=header)
    return {"heading": "PO Attainment Table:", "data": po_df}

def extract_po_evaluation_table(xls):
    """Extract PO evaluation table from CO Attainment sheet."""
    try:
        df = pd.read_excel(xls, sheet_name="CO Attainment", header=None)
        for i, row in df.iterrows():
            row_str = [str(cell).strip() if cell is not None else "" for cell in row]
            if any("PO Attainment" in cell for cell in row_str):
                if i + 4 < len(df):
                    header_row = df.iloc[i + 3]
                    data_row = df.iloc[i + 4]
                    # Extract PO headers and values, skipping CO code
                    header = [str(cell).strip() for cell in header_row[4:] if pd.notna(cell) and str(cell).strip()]
                    data = []
                    for cell in data_row[4:4+len(header)]:
                        val = str(cell).strip() if pd.notna(cell) else "-"
                        try:
                            num = float(val)
                            val = f"{num:.2f}"
                        except:
                            pass
                        data.append(val)
                    if len(header) > 0 and len(data) > 0:
                        return header, data
        return None
    except Exception as e:
        print(f"Error extracting PO evaluation table: {e}")
        return None

def extract_pso_evaluation_table(xls):
    """Extract the final PSO evaluation table (below 'Course' row) from PSO Attainment sheet."""
    try:
        df = pd.read_excel(xls, sheet_name="PSO Attainment", header=None)
        for i, row in df.iterrows():
            row_str = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cell == "Course" for cell in row_str):
                header_index = row_str.index("Course")
                header = row_str[header_index:]  # ['Course', 'PSO1', 'PSO2', 'PSO3']
                
                # Fetch next row (data row)
                if i + 1 < len(df):
                    data_row = df.iloc[i + 1]
                    data_values = [str(cell).strip() if pd.notna(cell) else "-" for cell in data_row[header_index:header_index+len(header)]]
                    
                    # Format numeric values
                    formatted_data = []
                    for val in data_values:
                        try:
                            num = float(val)
                            val = f"{num:.2f}"
                        except:
                            pass
                        formatted_data.append(val)

                    return header, formatted_data
        return None
    except Exception as e:
        print(f"Error extracting PSO evaluation table: {e}")
        return None
  
def read_excel(file_path):
    """Reads and returns CO‑PO, PO Attainment, and CO‑PSO mappings."""
    xls = pd.ExcelFile(file_path)
    co_po = extract_co_mapping(xls, "CO Attainment", "PO")
    po_att = extract_po_attainment(xls, "CO Attainment")
    co_pso = extract_co_mapping(xls, "PSO Attainment", "PSO")
    return co_po, po_att, co_pso



def create_word_report(co_po_map, po_att_map, co_pso_map, excel_file, output_path):
    """Builds and saves the .docx report."""
    doc = Document()

    # --- Logo and Title ---
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture("kbtimage.jpg", width=Inches(6))
    heading = doc.add_paragraph("TEACHER LEVEL ATTAINMENT REPORT/COURSE INFORMATION SHEET",style="Heading 1")  #style="Title"
    # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
  
    # Load workbook to extract faculty-related info
    wb = load_workbook(excel_file, data_only=True)
    faculty_info = {}
    keywords = ["Faculty Name", "A.Y", "Subject Name", "Sub.Code", "Sem", "Class"]

    # Search across all sheets
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            row = [str(cell).strip() if cell is not None else "" for cell in row]
            for i, cell in enumerate(row):
                if cell in keywords and i + 1 < len(row):
                    value = row[i + 1].strip() if isinstance(row[i + 1], str) else row[i + 1]
                    faculty_info[cell] = value

    # Fallback defaults if any field is missing
    for key in keywords:
        faculty_info.setdefault(key, "")


    # Create a table with 3 rows and 4 columns for faculty info
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    # Row 1
    table.rows[0].cells[0].text = "Faculty Name"
    table.rows[0].cells[1].text = str(faculty_info["Faculty Name"])
    table.rows[0].cells[2].text = "A.Y"
    table.rows[0].cells[3].text = str(faculty_info["A.Y"])

    # Row 2
    table.rows[1].cells[0].text = "Subject Name"
    table.rows[1].cells[1].text = str(faculty_info["Subject Name"])
    table.rows[1].cells[2].text = "Sub.Code"
    table.rows[1].cells[3].text = str(faculty_info["Sub.Code"])

    # Row 3
    table.rows[2].cells[0].text = "Sem"
    table.rows[2].cells[1].text = str(faculty_info["Sem"])
    table.rows[2].cells[2].text = "Class"
    table.rows[2].cells[3].text = str(faculty_info["Class"])


    # --- Teaching & Examination Scheme ---
    heading = doc.add_paragraph("Teaching & Examination Scheme:", style="Heading 1")

    # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Create table with 3 rows and 7 columns
    table = doc.add_table(rows=3, cols=7, style="Table Grid")

    # --- Merge and format "Teaching Scheme" ---
    teaching_cell = table.cell(0, 0).merge(table.cell(0, 2))
    para = teaching_cell.paragraphs[0]
    run = para.add_run("Teaching Scheme")
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Merge and format "Examination Scheme" ---
    exam_cell = table.cell(0, 3).merge(table.cell(0, 6))
    para = exam_cell.paragraphs[0]
    run = para.add_run("Examination Scheme")
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Second row headers
    headers = ["Theory", "Practical", "Tutorial", "In Sem", "End Sem", "PR", "Term Work"]
    for i, header in enumerate(headers):
        cell = table.cell(1, i)
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Third row values
    values = ["4 hrs/week", "-", "-", "✔", "✔", "-", "-"]
    for i, value in enumerate(values):
        cell = table.cell(2, i)
        cell.text = value
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # --- Delivery Method ---
    heading= doc.add_paragraph("Delivery Method:", style="Heading 1")
    # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Create table with 3 rows and 6 columns
    table = doc.add_table(rows=3, cols=6, style="Table Grid")

    # Merge the first row across all 6 columns for the "Delivery Method" title
    header_cell = table.cell(0, 0)
    header_cell.merge(table.cell(0, 5))
    header_cell.text = "Delivery Method"

    # Center align the merged cell text
    para = header_cell.paragraphs[0]
    para.alignment = 1  # 0=Left, 1=Center, 2=Right

    # Add method names to second row
    methods = ["Chalk & Talk", "ICT Tools", "Group Discussion", "Industrial Visit", "Expert Talk", "Virtual Lab"]
    for j, method in enumerate(methods):
        cell = table.cell(1, j)
        cell.text = method
        cell.paragraphs[0].alignment = 1  # center text

    # Add ticks to third row
    ticks = ["✔", "-", "✔", "-", "✔", "-"]
    for j, tick in enumerate(ticks):
        cell = table.cell(2, j)
        cell.text = tick
        cell.paragraphs[0].alignment = 1  # center text


    # Add the note 
    doc.add_paragraph("(* Kindly tick the methods conducted for this course. You may add any additional delivery method conducted in the above columns.)")

     # Program Outcomes (POs)
    heading = doc.add_paragraph("Program Outcomes:", style="Heading 1")
    # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    po_table = doc.add_table(rows=0, cols=2)
    po_table.style = 'Table Grid'

    # Add PO data rows
    po_data = [
        ("PO1", "Engineering knowledge: Apply the knowledge of mathematics, science, engineering fundamentals, and an engineering specialization to the solution of complex engineering problems."),
        ("PO2", "Problem analysis: Identify, formulate, review research literature, and analyze complex engineering problems reaching substantiated conclusions using first principles of mathematics, natural sciences, and engineering sciences."),
        ("PO3", "Design/development of solutions: Design solutions for complex engineering problems and design system components or processes that meet specified needs with appropriate consideration for public health and safety, and cultural, societal, and environmental considerations."),
        ("PO4", "Conduct investigations of complex problems: Use research-based knowledge and research methods including design of experiments, analysis and interpretation of data, and synthesis of the information to provide valid conclusions."),
        ("PO5", "Modern tool usage: Create, select, and apply appropriate techniques, resources, and modern engineering and IT tools including prediction and modeling to complex engineering activities with an understanding of the limitations."),
        ("PO6", "The engineer and society: Apply reasoning informed by contextual knowledge to assess societal, health, safety, legal and cultural issues and the consequent responsibilities relevant to the professional engineering practice."),
        ("PO7", "Environment and sustainability: Understand the impact of the professional engineering solutions in societal and environmental contexts, and demonstrate the knowledge of, and need for sustainable development."),
        ("PO8", "Ethics: Apply ethical principles and commit to professional ethics and responsibilities and norms of the engineering practice."),
        ("PO9", "Individual and team work: Function effectively as an individual, and as a member or leader in diverse teams, and in multidisciplinary settings."),
        ("PO10", "Communication: Communicate effectively on complex engineering activities with the engineering community and with society at large."),
        ("PO11", "Project management and finance: Demonstrate knowledge and understanding of the engineering and management principles and apply these to manage projects in multidisciplinary environments."),
        ("PO12", "Life-long learning: Recognize the need for, and have the preparation and ability to engage in independent and life-long learning in the broadest context of technological change.")
    ]

    for po_code, description in po_data:
        row_cells = po_table.add_row().cells
        # Set PO Code (bold and black)
        run = row_cells[0].paragraphs[0].add_run(po_code)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Set Description (normal style)
        row_cells[1].text = description

    # Adjust column widths
    for row in po_table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(5.5)


    #  Program Specific Outcomes (PSOs)
    heading = doc.add_paragraph("PSOs:", style="Heading 1")
    # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Create PSO table without header
    pso_table = doc.add_table(rows=0, cols=2)
    pso_table.style = "Table Grid"

    psos = [
        ("PSO1", "To demonstrate mathematical and Computer Engineering fundamentals"),
        ("PSO2", "To adapt modern computer tools and technologies to solve Computer Engineering Problems."),
        ("PSO3", "To apply software engineering practices and standards for project management")
    ]

    for pso_no, desc in psos:
        row_cells = pso_table.add_row().cells

        # Style first column as bold and black
        run = row_cells[0].paragraphs[0].add_run(pso_no)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        row_cells[1].text = desc

    # Adjust column widths
    for row in pso_table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(5.5)
    
   
    #  Course Outcomes (COs)
    
    doc.add_paragraph("(After completion of course, student will be able )")
    # CO Table without header row
    co_table = doc.add_table(rows=0, cols=2)
    co_table.style = "Table Grid"

    cos = [
        ("CO1", "Analyze needs and challenges for Data Science Big Data Analytics"),
        ("CO2", "Apply statistics for Big Data Analytics"),
        ("CO3", "Apply the lifecycle of Big Data analytics to real world problems"),
        ("CO4", "Implement Big Data Analytics using Python programming"),
        ("CO5", "Implement Big Data Analytics and model evaluation using algorithm."),
        ("CO6", "Design and implement Big Databases using the Hadoop ecosystem")
    ]

    for co_no, desc in cos:
        row_cells = co_table.add_row().cells

        # First column styled bold and black
        run = row_cells[0].paragraphs[0].add_run(co_no)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        row_cells[1].text = desc

    # Adjust column widths
    for row in co_table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(5.5)
    
      # --- 1) CO‑PO Mapping on New Page ---
    doc.add_page_break()
    heading= doc.add_paragraph(co_po_map["heading"], style="Heading 1")
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    df1 = co_po_map["data"]
    cols = len(df1.columns)
    rows = len(df1) + 2  # 2 header rows

    t1 = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    # Row 0: "COs" + merged "Program outcomes"
    co_cell = t1.cell(0, 0).merge(t1.cell(1, 0))  # ✅ Merge first column's top 2 rows
    co_cell.text = "COs"
    co_cell.paragraphs[0].runs[0].bold = True
    co_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Row 0: Merge PO headers
    po_header = t1.cell(0, 1)
    for i in range(2, cols):
        po_header = po_header.merge(t1.cell(0, i))
    po_header.text = "Program outcomes"
    po_header.paragraphs[0].runs[0].bold = True
    po_header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Row 1: PO1 to PO12 headers
    for j, col_name in enumerate(df1.columns[1:], start=1):  # Skip "CO" header
        cell = t1.cell(1, j)
        run = cell.paragraphs[0].add_run(str(col_name))
        run.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    for i, row in df1.iterrows():
        for j, val in enumerate(row):
            cell = t1.cell(i + 2, j)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

   # --- 3) CO‑PSO Mapping ---
    # doc.add_page_break()
    heading= doc.add_paragraph("Mapping of COs with PSOs:", style="Heading 1")
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    df3 = co_pso_map["data"]
    cols = len(df3.columns)
    rows = len(df3) + 2

    t3 = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    # Merge first column: "Course Outcomes"
    co_cell = t3.cell(0, 0).merge(t3.cell(1, 0))
    co_para = co_cell.paragraphs[0]
    co_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    co_run = co_para.add_run("Course Outcomes")
    co_run.bold = True

    # Merge "Program Specific Outcomes" across PSO1–PSO3
    header_cell = t3.cell(0, 1)
    for i in range(2, cols):
        header_cell = header_cell.merge(t3.cell(0, i))
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_run = header_para.add_run("Program Specific Outcomes")
    header_run.bold = True

    # Row 2: PSO1, PSO2, PSO3
    for j, col in enumerate(df3.columns[1:], start=1):
        cell = t3.cell(1, j)
        run = cell.paragraphs[0].add_run(str(col))
        run.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    for i, row in df3.iterrows():
        for j, val in enumerate(row):
            cell = t3.cell(i + 2, j)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adjust column widths and spacing
    # from docx.shared import Inches
    for row in t3.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Inches(1.5) if j == 0 else Inches(1.0)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0

    doc.add_paragraph()  # Optional spacing after table



    # --- Direct Assessment ---
    doc.add_paragraph()
    heading= doc.add_heading('A) Direct Assessment (90%)', level=2)
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading=doc.add_paragraph('B) External Assessment (80%): -')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    heading= doc.add_heading('Attainment Level Vs Target Value', level=1)
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)



    # Attainment Level Table
    direct_attain_table = doc.add_table(rows=1, cols=2)
    direct_attain_table.style = 'Table Grid'
    direct_attain_table.allow_autofit = True
    direct_attain_table.rows[0].cells[0].text = 'Attainment Level'
    direct_attain_table.rows[0].cells[1].text = 'Description'

    direct_levels = [
        ("1", "50% students scoring more than University Average marks or target value"),
        ("2", "60% students scoring more than University Average marks or target value"),
        ("3", "70% students scoring more than University Average marks or target value"),
    ]

    for level, desc in direct_levels:
        row = direct_attain_table.add_row().cells
        row[0].text = level
        row[1].text = desc

      # Target Value Table for Direct Assessment
    def extract_target_value(wb):
        """Extract target value from 'University TH Marks' sheet."""
        try:
            sheet = wb["University TH Marks"]
            for row in sheet.iter_rows(values_only=True):
                row_str = [str(cell).strip() if cell is not None else "" for cell in row]
                if any("number of students scoring >=" in cell.lower() for cell in row_str):
                    for cell in row:
                        if isinstance(cell, (int, float)):
                            return str(int(cell)) if isinstance(cell, int) else f"{cell:.2f}"
            return "-"
        except Exception as e:
            print(f"Error extracting target value: {e}")
            return "-"
   
   
    doc.add_paragraph()
    heading= doc.add_heading('Set Target Value', level=1)
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    direct_target_table = doc.add_table(rows=2, cols=3)
    direct_target_table.style = 'Table Grid'
    direct_target_table.rows[0].cells[0].text = 'Theory'
    direct_target_table.rows[0].cells[1].text = 'PR'
    direct_target_table.rows[0].cells[2].text = 'Term Work'

    # direct_target_table.rows[1].cells[0].text = '49'
    target_value = extract_target_value(wb)
    direct_target_table.rows[1].cells[0].text = target_value

    direct_target_table.rows[1].cells[1].text = '-'
    direct_target_table.rows[1].cells[2].text = '-'

    # --- Internal Assessment (20%) ---
    doc.add_paragraph()
    heading= doc.add_heading('C) Internal Assessment (20%): -',level=2)
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    heading= doc.add_paragraph('Attainment Level Vs Target Value', style='Heading 2')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    internal_attain_table = doc.add_table(rows=1, cols=2)
    internal_attain_table.style = 'Table Grid'
    internal_attain_table.rows[0].cells[0].text = 'Attainment Level'
    internal_attain_table.rows[0].cells[1].text = 'Description'

    internal_levels = [
        ("1", "60% students scoring more than 60% of maximum marks"),
        ("2", "70% students scoring more than 60% of maximum marks"),
        ("3", "80% students scoring more than 60% of maximum marks"),
    ]

    for level, desc in internal_levels:
        row = internal_attain_table.add_row().cells
        row[0].text = level
        row[1].text = desc

    # Target Value Table for Internal Assessment
    doc.add_paragraph()
    heading= doc.add_heading('Set Target Value', level=2)
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    internal_target_table = doc.add_table(rows=2, cols=2)
    internal_target_table.style = 'Table Grid'
    internal_target_table.rows[0].cells[0].text = 'Unit Test'
    internal_target_table.rows[0].cells[1].text = 'Term Work'
    internal_target_table.rows[1].cells[0].text = '60%'
    internal_target_table.rows[1].cells[1].text = '60%'
       
    # === External Assessment ===
    sheet_external = wb['University TH Marks']
    external_scores = []

    for row in sheet_external.iter_rows(values_only=True):
        for cell in row:
            if cell and isinstance(cell, str) and "Percentage  of Students scoring >= Target Value" in cell:
                extracted_values = [c for c in row if isinstance(c, (int, float, str))]
                for val in extracted_values:
                    try:
                        raw = float(str(val))
                        percentage = raw * 100 if raw <= 1 else raw
                        external_scores = [percentage] * 6  # Use same value for CO1 to CO6
                        break
                    except:
                        continue
                break
        if external_scores:
            break

    external_scores = external_scores[:6] + [0.0] * (6 - len(external_scores))  # Fallback if less than 6 values

    # doc.add_paragraph("\nExternal Assessment:")
    heading= doc.add_paragraph('External Assessment:',style='Heading 1')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("Theory (%)")
    table_ext = doc.add_table(rows=2, cols=6)
    table_ext.style = 'Table Grid'
    for i in range(6):
        table_ext.cell(0, i).text = f"CO{i+1}"
        table_ext.cell(1, i).text = f"{external_scores[i]:.2f}"  # No % sign
 

     # Internal Assessment
    sheet_ia = wb['Internal Assessment marks']
    internal_scores = []
    for row in sheet_ia.iter_rows(values_only=True):
        for cell in row:
            if cell and isinstance(cell, str) and "CO Attainment" in cell:
                internal_scores = [c for c in row if isinstance(c, (int, float))]
                break
        if internal_scores:
            break

    internal_scores = internal_scores[:6] + [1.0] * (6 - len(internal_scores))
    # doc.add_paragraph("\nInternal Assessment:")
    heading= doc.add_paragraph('Internal Assessment:', style='Heading 1')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("Unit Test (%)")
    table_int = doc.add_table(rows=2, cols=6)
    table_int.style = 'Table Grid'
    for i in range(6):
        table_int.cell(0, i).text = f"CO{i+1}"
        table_int.cell(1, i).text = f"{internal_scores[i] * 100:.2f}"

    # Indirect Assessment
    sheet_exit = wb['Course Exit survey']
    exit_scores = []
    for row in sheet_exit.iter_rows(values_only=True):
        if row and any("Attainment level in percentage" in str(cell) for cell in row):
            exit_scores = [c for c in row if isinstance(c, (int, float))]
            break

    exit_scores = exit_scores[:6] + [0.0] * (6 - len(exit_scores))
    # doc.add_paragraph("\nIndirect Assessment:")
    heading= doc.add_paragraph('Indirect Assessment:', style='Heading 1')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("Course Exit Survey (%)")
    table_exit = doc.add_table(rows=2, cols=6)
    table_exit.style = 'Table Grid'
    for i in range(6):
        table_exit.cell(0, i).text = f"CO{i+1}"
        table_exit.cell(1, i).text = f"{exit_scores[i]:.2f}"

    # External Attainment
    external_attainment = None
    for row in wb['University TH Marks'].iter_rows(values_only=True):
        if row and any("Attainment Level" in str(cell) for cell in row):
            for val in reversed(row):
                if isinstance(val, (int, float)):
                    external_attainment = val
                    break
            break

    # doc.add_paragraph("\nCO Attainment by External Assessment")
    heading= doc.add_paragraph('CO Attainment by External Assessment', style='Heading 1')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table_ext = doc.add_table(rows=3, cols=2)
    table_ext.style = 'Table Grid'
    table_ext.cell(0, 0).text = "External Assessment Tool"
    table_ext.cell(0, 1).text = "CO Attainment"
    table_ext.cell(1, 0).text = "TH Exam"
    table_ext.cell(1, 1).text = f"{external_attainment}"
    table_ext.cell(2, 0).text = "Average"
    table_ext.cell(2, 1).text = f"{external_attainment}"

    # Internal Attainment
    internal_attainment = None
    for row in wb['Internal Assessment marks'].iter_rows(values_only=True):
        if row and any("Target Level" in str(cell) for cell in row):
            for val in reversed(row):
                if isinstance(val, (int, float)):
                    internal_attainment = val
                    break
            break

    heading= doc.add_paragraph("CO Attainment by Internal Assessment",style='Heading 1')
     # Set heading font color to black
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table_int2 = doc.add_table(rows=3, cols=2)
    table_int2.style = 'Table Grid'
    table_int2.cell(0, 0).text = "Internal Assessment Tool"
    table_int2.cell(0, 1).text = "CO Attainment"
    table_int2.cell(1, 0).text = "Unit Test"
    table_int2.cell(1, 1).text = f"{internal_attainment}"
    table_int2.cell(2, 0).text = "Average"
    table_int2.cell(2, 1).text = f"{internal_attainment}"

    # # Final Attainment Calculation
    # university_attainment = external_attainment or 3
    # internal_attainment = internal_attainment or 2
    # final_attainment = 0.8 * university_attainment + 0.2 * internal_attainment

    # # doc.add_paragraph("\nFinal CO Attainment Calculation:")
    # heading=doc.add_paragraph('Final CO Attainment Calculation:', style='Heading 2')
    # # Set heading font color to black
    # run = heading.runs[0]
    # run.font.color.rgb = RGBColor(0, 0, 0)

    # doc.add_paragraph("Direct CO attainment is computed as:")
    # doc.add_paragraph("= 0.8 × CO attainment level in university examination")
    # doc.add_paragraph("+ 0.2 × CO attainment level in internal assessment")
    # doc.add_paragraph(f"= 0.8 × {university_attainment} + 0.2 × {internal_attainment}")
    # doc.add_paragraph(f"= {final_attainment:.2f}")

    # === CO Attainment by Course Exit Survey (based on Final Attainment Level) ===
    exit_survey_sheet = wb['Course Exit survey']
    exit_survey_attainment = None

    # Find "Final attainment level" and extract the numeric value from the row
    for row in exit_survey_sheet.iter_rows(values_only=True):
        if row and any("Final attainment level" in str(cell) for cell in row):
            for val in reversed(row):
                if isinstance(val, (int, float)):
                    exit_survey_attainment = val
                    break
            break

    # === CO Attainment by Course Exit Survey Section ===
    heading = doc.add_paragraph("CO Attainment by Course Exit Survey", style='Heading 1')
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    exit_survey_table = doc.add_table(rows=2, cols=2)
    exit_survey_table.style = 'Table Grid'

    # Table Headers
    exit_survey_table.cell(0, 0).text = "Indirect Assessment Tool"
    exit_survey_table.cell(0, 1).text = "CO Attainment"

    # Fill Data Row
    attainment_value = exit_survey_attainment if exit_survey_attainment is not None else "-"
    exit_survey_table.cell(1, 0).text = "Course Exit Survey"
    exit_survey_table.cell(1, 1).text = str(attainment_value)


    # # === Final CO Attainment Calculation ===
    # # Replace these with actual calculated values if available
    # co_attainment_direct = 2.8
    # co_attainment_indirect = 3.0

    # final_co_attainment = 0.9 * co_attainment_direct + 0.1 * co_attainment_indirect

    # doc.add_paragraph("\nOverall CO attainment is then computed as:")
    # doc.add_paragraph("= 0.9 × CO attainment level in Direct CO attainment")
    # doc.add_paragraph("+ 0.1 × CO attainment level in Indirect CO attainment")
    # doc.add_paragraph(f"\nFinal CO Attainment = {final_co_attainment:.2f}")


    # === Result of Evaluation of PO's Table ===
    try:
        xls = pd.ExcelFile(excel_file)
        po_eval = extract_po_evaluation_table(xls)
        if po_eval:
            doc.add_paragraph()
            heading= doc.add_paragraph("Result of Evaluation of PO's Table", style='Heading 1')
             # Set heading font color to black
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
            header, data = po_eval
            if len(header) > 0 and len(data) > 0:
                # Create table with proper dimensions
                max_cols = max(len(header), len(data))
                table = doc.add_table(rows=2, cols=max_cols, style="Table Grid")
                
                # Fill header row
                for j, h in enumerate(header):
                    if j < max_cols:
                        table.cell(0, j).text = str(h)
                
                # Fill data row
                for j, d in enumerate(data):
                    if j < max_cols:
                        table.cell(1, j).text = str(d)
            doc.add_paragraph()
    except Exception as e:
        print(f"Error creating PO evaluation table: {e}")

    # === Result of Evaluation of PSO's Table ===
    doc.add_page_break()
    try:
        pso_eval = extract_pso_evaluation_table(xls)
        if pso_eval:
            doc.add_paragraph("Result of Evaluation of PSO's Table", style='Heading 1')
            header, data = pso_eval
            if len(header) > 0 and len(data) > 0:
                table = doc.add_table(rows=2, cols=len(header), style="Table Grid")
                
                # Header row
                for j, h in enumerate(header):
                    table.cell(0, j).text = str(h)
                
                # Data row
                for j, d in enumerate(data):
                    table.cell(1, j).text = str(d)

            doc.add_paragraph()
    except Exception as e:
        print(f"Error creating PSO evaluation table: {e}")


    # === Remarks and Action Plan Section ===
    doc.add_paragraph("\nRemarks:")

    try:
        final_co_attainment = float(exit_survey_attainment)
    except:
        final_co_attainment = 0  # fallback if value is missing or invalid

    if final_co_attainment >= 3:
        remark_text = "Target is Achieved."
        observation_text = "Students have performed well and met the expected outcomes."
        action_plan_text = "Continue with the current teaching strategies and maintain performance."
    else:
        remark_text = "Target is not Achieved."
        observation_text = "Need to take more practice of algorithms."
        action_plan_text = "University questions to be solved for better understanding."

    # Add Remarks Content
    para1 = doc.add_paragraph()
    para1.add_run("Remark: ").bold = True
    para1.add_run(remark_text)

    para2 = doc.add_paragraph()
    para2.add_run("Observation: ").bold = True
    para2.add_run(observation_text)

    para3 = doc.add_paragraph()
    para3.add_run("Action Plan: ").bold = True
    para3.add_run(action_plan_text)

    # === Target Set for A.Y. 2024-25 Section ===
    from datetime import datetime

    # Get current year and next year
    current_year = datetime.now().year
    next_year = current_year + 1

    # Format Academic Year (e.g., "2024-25")
    ay_string = f"{current_year}-{str(next_year)[-2:]}"  # e.g., "2024-25"

    # Add to Word document
    para4 = doc.add_paragraph()
    para4.add_run(f"Target Set for A.Y. {ay_string}:").bold = True



    target_ay_headers = ["Theory", "Oral / Practical", "Term Work"]
    target_ay_values = ["50", "NA", "NA"]

    ay_target_table = doc.add_table(rows=2, cols=3)
    ay_target_table.style = 'Table Grid'

    # Fill header
    for col_index in range(3):
        ay_target_table.cell(0, col_index).text = target_ay_headers[col_index]

    # Fill values
    for col_index in range(3):
        ay_target_table.cell(1, col_index).text = target_ay_values[col_index]

    
    
    # Save to disk
    doc.save(output_path)

    

@app.route('/upload', methods=["POST"])
def upload():
    f = request.files.get('file')
    if not f:
        return jsonify(status="error", message="No file uploaded.")
    name = secure_filename(f.filename)
    in_path = os.path.join(UPLOAD_FOLDER, name)
    f.save(in_path)

    try:
        co_po_map, po_att_map, co_pso_map = read_excel(in_path)
        out_name = f"report_{os.path.splitext(name)[0]}.docx"
        out_path = os.path.join(REPORT_FOLDER, out_name)
        create_word_report(co_po_map, po_att_map, co_pso_map, in_path, out_path)
        return jsonify(status="success", filename=out_name)
    except Exception as e:
        return jsonify(status="error", message=str(e))


@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(REPORT_FOLDER, filename, as_attachment=True)


@app.route('/')
def index_page():
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
